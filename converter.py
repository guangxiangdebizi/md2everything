#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown 转换工具 - 简化版
支持 HTML 和 DOCX 导出，PDF 通过浏览器打印实现
"""

import re
import markdown
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from bs4 import BeautifulSoup


class MarkdownConverter:
    """Markdown 转换器"""
    
    def __init__(self):
        self.md = markdown.Markdown(
            extensions=[
                'extra',
                'codehilite',
                'tables',
                'toc',
                'fenced_code',
                'attr_list',
            ],
            extension_configs={
                'codehilite': {
                    'linenums': False,
                    'guess_lang': False,
                },
            }
        )
    
    def _get_html_template(self, content, title="Document"):
        """生成完整的 HTML 文档"""
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        
        @media print {{
            body {{
                padding: 0;
                background: white;
            }}
            .no-print {{
                display: none;
            }}
        }}
        
        body {{
            font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
            line-height: 1.8;
            color: #333;
            font-size: 11pt;
            max-width: 1200px;
            margin: 0 auto;
            padding: 40px 20px;
            background: #f8f9fa;
        }}
        
        .container {{
            background: white;
            padding: 40px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            border-radius: 8px;
        }}
        
        .toolbar {{
            position: fixed;
            top: 20px;
            right: 20px;
            background: white;
            padding: 15px;
            border-radius: 8px;
            box-shadow: 0 4px 12px rgba(0,0,0,0.15);
            z-index: 1000;
        }}
        
        .btn {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 6px;
            cursor: pointer;
            font-size: 14px;
            font-weight: 600;
            transition: transform 0.2s;
        }}
        
        .btn:hover {{
            transform: translateY(-2px);
        }}
        
        h1 {{
            font-size: 2.2em;
            color: #667eea;
            border-bottom: 3px solid #667eea;
            padding-bottom: 0.3em;
            margin-top: 1.5em;
            margin-bottom: 0.8em;
            page-break-after: avoid;
        }}
        
        h2 {{
            font-size: 1.8em;
            color: #495057;
            border-bottom: 2px solid #e9ecef;
            padding-bottom: 0.3em;
            margin-top: 1.3em;
            margin-bottom: 0.6em;
            page-break-after: avoid;
        }}
        
        h3 {{
            font-size: 1.4em;
            color: #6c757d;
            margin-top: 1.2em;
            margin-bottom: 0.5em;
            page-break-after: avoid;
        }}
        
        h4 {{
            font-size: 1.1em;
            color: #868e96;
            margin-top: 1em;
            margin-bottom: 0.4em;
        }}
        
        p {{
            margin-bottom: 0.8em;
            text-align: justify;
        }}
        
        ul, ol {{
            margin-bottom: 1em;
            padding-left: 2em;
        }}
        
        li {{
            margin-bottom: 0.4em;
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1.5em 0;
            page-break-inside: avoid;
        }}
        
        table th {{
            background: #667eea;
            color: white;
            padding: 10px;
            text-align: left;
            font-weight: 600;
            border: 1px solid #5568d3;
        }}
        
        table td {{
            padding: 8px 10px;
            border: 1px solid #e9ecef;
        }}
        
        table tr:nth-child(even) {{
            background: #f8f9fa;
        }}
        
        code {{
            background: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Consolas', 'Monaco', monospace;
            font-size: 0.9em;
            color: #e83e8c;
        }}
        
        pre {{
            background: #f8f9fa;
            border: 1px solid #e9ecef;
            border-left: 4px solid #667eea;
            padding: 15px;
            border-radius: 4px;
            overflow-x: auto;
            margin: 1.5em 0;
            page-break-inside: avoid;
        }}
        
        pre code {{
            background: transparent;
            padding: 0;
            color: #333;
            font-size: 0.85em;
            line-height: 1.5;
        }}
        
        blockquote {{
            border-left: 4px solid #667eea;
            padding-left: 20px;
            margin: 1.5em 0;
            color: #6c757d;
            font-style: italic;
        }}
        
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1.5em auto;
            page-break-inside: avoid;
        }}
        
        hr {{
            border: none;
            border-top: 2px solid #e9ecef;
            margin: 2em 0;
        }}
        
        .mermaid-note {{
            text-align: center;
            padding: 20px;
            margin: 1.5em 0;
            background: #fff3cd;
            border: 2px solid #ffc107;
            border-radius: 8px;
            color: #856404;
            font-weight: 600;
        }}
    </style>
</head>
<body>
    <div class="toolbar no-print">
        <button class="btn" onclick="window.print()">🖨️ 打印/保存为PDF</button>
    </div>
    
    <div class="container">
        {content}
    </div>
    
    <script>
        // 键盘快捷键：Ctrl+P 打印
        document.addEventListener('keydown', function(e) {{
            if ((e.ctrlKey || e.metaKey) && e.key === 'p') {{
                e.preventDefault();
                window.print();
            }}
        }});
    </script>
</body>
</html>
"""
    
    def _process_mermaid(self, html_content):
        """处理 Mermaid 代码块"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for code in soup.find_all('code', class_=re.compile(r'language-mermaid')):
            note = soup.new_tag('div', **{'class': 'mermaid-note'})
            note.string = '📊 Mermaid 图表（在前端版本 index.html 中可查看完整图表）'
            
            if code.parent and code.parent.name == 'pre':
                code.parent.replace_with(note)
        
        return str(soup)
    
    def to_html(self, md_content, title="Document"):
        """转换为 HTML"""
        html_body = self.md.convert(md_content)
        html_body = self._process_mermaid(html_body)
        html_full = self._get_html_template(html_body, title)
        return html_full
    
    def to_docx(self, md_content):
        """转换为 DOCX（返回字节流）"""
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.font.size = Pt(11)
        
        # 解析 Markdown
        html = self.md.convert(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        
        # 处理元素
        self._process_docx_element(doc, soup)
        
        # 保存到字节流
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes
    
    def _process_docx_element(self, doc, element):
        """递归处理 HTML 元素转 DOCX"""
        for child in element.children:
            if not hasattr(child, 'name') or not child.name:
                continue
            
            if child.name in ['h1', 'h2', 'h3', 'h4']:
                level = int(child.name[1])
                heading = doc.add_heading(child.get_text().strip(), level=level)
                
                if level == 1:
                    heading.runs[0].font.size = Pt(24)
                    heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
                elif level == 2:
                    heading.runs[0].font.size = Pt(20)
                    heading.runs[0].font.color.rgb = RGBColor(73, 80, 87)
                elif level == 3:
                    heading.runs[0].font.size = Pt(16)
                    heading.runs[0].font.color.rgb = RGBColor(108, 117, 125)
            
            elif child.name == 'p':
                text = child.get_text().strip()
                if text:
                    doc.add_paragraph(text)
            
            elif child.name in ['ul', 'ol']:
                for li in child.find_all('li', recursive=False):
                    para = doc.add_paragraph(li.get_text().strip())
                    para.style = 'List Bullet' if child.name == 'ul' else 'List Number'
            
            elif child.name == 'table':
                self._add_docx_table(doc, child)
            
            elif child.name == 'pre':
                code = child.find('code')
                if code:
                    code_text = code.get_text()
                    classes = code.get('class', [])
                    
                    # 检查是否为 Mermaid
                    if any('mermaid' in str(c) for c in classes):
                        para = doc.add_paragraph('📊 [Mermaid 图表 - 请在 HTML/前端版本查看]')
                        para.runs[0].font.color.rgb = RGBColor(102, 126, 234)
                        para.runs[0].font.bold = True
                    else:
                        para = doc.add_paragraph(code_text)
                        para.paragraph_format.left_indent = Inches(0.5)
                        para.runs[0].font.name = 'Consolas'
                        para.runs[0].font.size = Pt(9)
            
            elif child.name == 'blockquote':
                para = doc.add_paragraph(child.get_text().strip())
                para.paragraph_format.left_indent = Inches(0.5)
                para.runs[0].font.italic = True
                para.runs[0].font.color.rgb = RGBColor(108, 117, 125)
            
            elif child.name == 'hr':
                doc.add_paragraph('─' * 50)
    
    def _add_docx_table(self, doc, table_element):
        """添加表格到 DOCX"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        first_row = rows[0]
        cols = len(first_row.find_all(['th', 'td']))
        
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Light Grid Accent 1'
        
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for col_idx, cell in enumerate(cells):
                table_cell = table.rows[row_idx].cells[col_idx]
                table_cell.text = cell.get_text().strip()
                
                if row_idx == 0:
                    for paragraph in table_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    
                    shading_elm = table_cell._element.get_or_add_tcPr()
                    shading = shading_elm.get_or_add_shd()
                    shading.set(qn('w:fill'), '667EEA')


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 3:
        print("使用方法:")
        print("  python converter.py <input.md> <output.html>")
        print("  python converter.py <input.md> <output.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    converter = MarkdownConverter()
    ext = output_file.lower().split('.')[-1]
    
    if ext in ['html', 'htm']:
        html = converter.to_html(md_content, title=input_file)
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html)
        print(f"✓ HTML 已生成: {output_file}")
        print(f"💡 提示: 打开 HTML 文件，按 Ctrl+P 或点击按钮即可保存为 PDF")
    
    elif ext in ['docx', 'doc']:
        docx_bytes = converter.to_docx(md_content)
        with open(output_file, 'wb') as f:
            f.write(docx_bytes.read())
        print(f"✓ Word 已生成: {output_file}")
    
    else:
        print(f"不支持的格式: {ext}")
        print("支持的格式: html, docx")
        sys.exit(1)
